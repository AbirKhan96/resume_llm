import transformers
import torch
import os
model_id = "meta-llama/Meta-Llama-3.1-8B-Instruct"
os.environ["TOKENIZERS_PARALLELISM"] = "false"
pipeline = transformers.pipeline(
    "text-generation",
    model=model_id,
    model_kwargs={"torch_dtype": torch.bfloat16},
    device_map="auto",
)

import pandas as pd
from pathlib import Path
import sys
import docx

from docx import Document

from pathlib import Path
resume_texts={}
domains=[f for fold in Path('/home/aiml/hemanth/jackmack/llms/ner/Resume_bulk').iterdir() for f in fold.iterdir()]
# print(domains)
domain_name=[str(d.name) for d in domains]
for row in domains:

  print(row)
  document = Document(row)#Document(row.Path)

  doc_text = '\n'.join([paragraph.text for paragraph in document.paragraphs])
  # print(doc_text)
  resume_texts[str(row)]=doc_text
  
entity_prefixes = {
    "Name":[],
    "Email": [],
    "Gender": [],
    "Phone_Number": [],
    "Nationality": [],
    "LinkedIn": [],
    "Degrees": [],
    "Designation": [],
    "Last_Company": [],
    "Skills": [],
    "Years_of_Experience":[]
}


import pandas as pd

# Initialize a list to hold data for each resume
resume_data = []

for path, text in resume_texts.items():
    # Step 1: Extract years of experience
    experience_prompt = [
        {"role": "system", "content": "Calculate the total number of years and months of professional experience mentioned in the text, using the current year, 2024, for ongoing roles. Only periods of active employment should be included, excluding any gaps between jobs. Provide the result as a single numerical value in years, including months as a decimal, e.g., 'X.Y years'."},
        {"role": "user", "content": f"{text}"},
        {"role": "user", "content": "Calculate the total professional experience from the provided job durations, excluding gaps. Use 2024 for ongoing roles. Provide the result in 'X.Y years' format, with no additional explanations or details."}
    ]
    
    outputs = pipeline(
        experience_prompt,
        max_new_tokens=250,
        temperature=0.1,  # Adjust temperature for less randomness
        top_k=50,         # Consider top 50 tokens
        top_p=0.9         # Nucleus sampling
    )
    
    years_promt = outputs[0]["generated_text"][-1]['content']
    
    experience_prompt = [
        {"role": "user", "content": f"{years_promt}"},
        {"role": "user", "content": "Total professional experience: ? without any additional information or explanation "}
    ]
    
    outputs = pipeline(
        experience_prompt,
        max_new_tokens=10,
        temperature=0.1,  # Adjust temperature for less randomness
        top_k=50,         # Consider top 50 tokens
        top_p=0.9         # Nucleus sampling
    )
    
    yoe = outputs[0]["generated_text"][-1]['content']
    
    # Step 2: Extract skills
    messages = [
        {"role": "system", "content": "List only the skills that are tangible and can be directly applied. Avoid listing activities, roles, or abstract concepts."},
        {"role": "user", "content": f"{text}"},
        {"role": "user", "content": "Please provide a list of tangible skills only, excluding roles, activities, and responsibilities."}
    ]
    
    outputs = pipeline(
        messages,
        max_new_tokens=256,
        temperature=0.3,  # Adjust temperature for less randomness
        top_k=50,         # Consider top 50 tokens
        top_p=0.9         # Nucleus sampling
    )
    
    skills = outputs[0]["generated_text"][-1]['content']
    
    # Step 3: Extract other entities
    prompts = {
        "Name": "Provide the name of the person.",
        "Gender": "Provide the gender of the person.",
        "Education": "Provide the education qualifications of the person.",
        "Last_Company": "Provide the name of the last company the person worked at.",
        "Email": "Provide the email address of the person.",
        "LinkedIn": "Provide the LinkedIn profile of the person.",
        "Phone_Number": "Provide the phone number of the person.",
        "Nationality": "Provide the Nationality of the person.",
        "Designation": "Provide the Last Designation of the person.",
        "Degrees": "Provide the Degrees of the person."
    }
    
    # Initialize an empty dictionary to store the responses
    responses = {}

    # Iterate over each type of information
    for key, prompt in prompts.items():
        messages = [
            {"role": "system", "content": "Provide direct answers without any explanation."},
            {"role": "user", "content": f"{text}"},
            {"role": "user", "content": prompt}
        ]
        
        # Get the output for the current question
        output = pipeline(messages, max_new_tokens=100, temperature=0.5, top_k=50, top_p=0.9)
        
        # Store the response in the dictionary
        responses[key] = output[0]["generated_text"][-1]['content']
    
    # Add the years of experience and skills to the responses dictionary
    responses['Years_of_Experience'] = yoe
    responses['Skills'] = skills
    
    # Add the responses dictionary to the resume_data list
    resume_data.append(responses)
    # break
# Convert the list of dictionaries into a DataFrame
df = pd.DataFrame(resume_data)

# Display the DataFrame
print(df)

df.to_csv('/home/aiml/hemanth/jackmack/llms/ner/resumes_data_extraction.csv',index=False)
